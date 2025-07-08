import io
import os
import pdfplumber
from docx import Document
import pytesseract
from PIL import Image
import pandas as pd
from pdf2image import convert_from_path
import fitz  # PyMuPDF

# Set Tesseract path
pytesseract.pytesseract.tesseract_cmd = r"C:/Users/USER/AppData/Local/Programs/Tesseract-OCR/tesseract.exe"

def extract_text_from_pdf_with_ocr(file_path):
    """Extract text from PDF including OCR for images and proper table handling"""
    text = ""
    
    # First try regular text extraction with table support
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # Extract regular text
                page_text = page.extract_text() or ""
                text += f"\n--- Page {page_num + 1} ---\n"
                text += page_text
                
                # Extract tables
                tables = page.extract_tables()
                if tables:
                    text += "\n[TABLES FOUND]\n"
                    for table_num, table in enumerate(tables):
                        text += f"\nTable {table_num + 1}:\n"
                        # Convert table to pandas DataFrame for better formatting
                        if table:
                            df = pd.DataFrame(table[1:], columns=table[0] if table[0] else None)
                            text += df.to_string(index=False) + "\n"
    except Exception as e:
        pass
    
    # If no text found, try OCR
    if len(text.strip()) < 50:  # Threshold for "no meaningful text"
        try:
            # Convert PDF to images
            images = convert_from_path(file_path, dpi=300)
            text = ""
            for page_num, image in enumerate(images):
                text += f"\n--- Page {page_num + 1} (OCR) ---\n"
                # Use OCR to extract text from image
                ocr_text = pytesseract.image_to_string(image, lang='eng')
                text += ocr_text
        except Exception as e:
            pass
    
    return text

def extract_images_and_ocr_from_pdf(file_path):
    """Extract images from PDF and apply OCR"""
    text = ""
    try:
        pdf_document = fitz.open(file_path)
        
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            image_list = page.get_images()
            
            if image_list:
                text += f"\n--- Images from Page {page_num + 1} ---\n"
                
            for img_index, img in enumerate(image_list):
                # Get image data
                xref = img[0]
                pix = fitz.Pixmap(pdf_document, xref)
                
                if pix.n - pix.alpha < 4:  # GRAY or RGB
                    # Convert to PIL Image
                    img_data = pix.tobytes("ppm")
                    pil_image = Image.open(io.BytesIO(img_data))
                    
                    # Apply OCR
                    ocr_text = pytesseract.image_to_string(pil_image)
                    if ocr_text.strip():
                        text += f"\nImage {img_index + 1} text:\n{ocr_text}\n"
                
                pix = None
        
        pdf_document.close()
    except Exception as e:
        pass
    
    return text

def extract_text_from_docx(file_path):
    """Extract text from DOCX including tables"""
    text = ""
    try:
        doc = Document(file_path)
        
        # Extract paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extract tables
        if doc.tables:
            text += "\n[TABLES FOUND]\n"
            for table_num, table in enumerate(doc.tables):
                text += f"\nTable {table_num + 1}:\n"
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    df = pd.DataFrame(table_data[1:], columns=table_data[0] if table_data else None)
                    text += df.to_string(index=False) + "\n"
                    
    except Exception as e:
        pass
    
    return text

def extract_text_from_txt(file_path):
    """Extract text from TXT files"""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        return ""

def extract_text_from_image(file_path):
    """Extract text from image files using OCR"""
    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image, lang='eng')
        return text
    except Exception as e:
        return ""

def process_resumes():
    """Main function to extract text from various file types"""
    input_folder = "aadhar"
    output_folder = "aadhar_txt"
    
    # Statistics tracking
    stats = {
        'total_resumes': 0,
        'processed_successfully': 0,
        'failed_processing': 0,
        'errors': []
    }
    
    os.makedirs(output_folder, exist_ok=True)
    
    supported_extensions = {'.pdf', '.docx', '.txt', '.png', '.jpg', '.jpeg', '.tiff', '.bmp'}
    
    # Count total resume files
    all_files = os.listdir(input_folder)
    for filename in all_files:
        _, ext = os.path.splitext(filename)
        if ext.lower() in supported_extensions:
            stats['total_resumes'] += 1
    
    print("Processing...")
    
    for filename in all_files:
        file_path = os.path.join(input_folder, filename)
        base_name, ext = os.path.splitext(filename)
        txt_filename = base_name + ".txt"
        txt_path = os.path.join(output_folder, txt_filename)
        
        try:
            ext = ext.lower()
            text = ""
            
            if ext not in supported_extensions:
                continue
            
            if ext == ".pdf":
                text = extract_text_from_pdf_with_ocr(file_path)
                # Also try to extract text from images within PDF
                image_text = extract_images_and_ocr_from_pdf(file_path)
                if image_text:
                    text += "\n" + image_text
                    
            elif ext == ".docx":
                text = extract_text_from_docx(file_path)
                
            elif ext == ".txt":
                text = extract_text_from_txt(file_path)
                
            elif ext in {'.png', '.jpg', '.jpeg', '.tiff', '.bmp'}:
                text = extract_text_from_image(file_path)
            
            # Save extracted text
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(text)
            
            stats['processed_successfully'] += 1
            
        except Exception as e:
            error_msg = f"Failed to process {filename}: {str(e)}"
            stats['errors'].append(error_msg)
            stats['failed_processing'] += 1
    
    print("Process completed")
    
    return stats

if __name__ == "__main__":
    process_resumes()